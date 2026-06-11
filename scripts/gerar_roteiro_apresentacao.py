from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


RAIZ = Path(__file__).resolve().parents[1]
SAIDA = RAIZ / "docs" / "Roteiro_Apresentacao_Appono_12min.docx"

MARROM = "402820"
CARAMELO = "A85E32"
DOURADO = "D6A04C"
CREME = "F7F0E7"
CREME_ESCURO = "E9D9C4"
CINZA = "5B5652"
BRANCO = "FFFFFF"
VERMELHO = "9B1C1C"


def definir_fonte(run, nome="Aptos", tamanho=10.5, cor=MARROM, negrito=False, italico=False):
    run.font.name = nome
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), nome)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), nome)
    run.font.size = Pt(tamanho)
    run.font.color.rgb = RGBColor.from_string(cor)
    run.bold = negrito
    run.italic = italico


def sombrear_celula(celula, cor):
    propriedades = celula._tc.get_or_add_tcPr()
    sombreado = propriedades.find(qn("w:shd"))
    if sombreado is None:
        sombreado = OxmlElement("w:shd")
        propriedades.append(sombreado)
    sombreado.set(qn("w:fill"), cor)


def margens_celula(celula, top=100, start=140, bottom=100, end=140):
    propriedades = celula._tc.get_or_add_tcPr()
    tc_mar = propriedades.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        propriedades.append(tc_mar)
    for lado, valor in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        item = tc_mar.find(qn(f"w:{lado}"))
        if item is None:
            item = OxmlElement(f"w:{lado}")
            tc_mar.append(item)
        item.set(qn("w:w"), str(valor))
        item.set(qn("w:type"), "dxa")


def bordas_tabela(tabela, cor="D7C3AA", tamanho="6"):
    propriedades = tabela._tbl.tblPr
    bordas = propriedades.first_child_found_in("w:tblBorders")
    if bordas is None:
        bordas = OxmlElement("w:tblBorders")
        propriedades.append(bordas)
    for nome in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borda = bordas.find(qn(f"w:{nome}"))
        if borda is None:
            borda = OxmlElement(f"w:{nome}")
            bordas.append(borda)
        borda.set(qn("w:val"), "single")
        borda.set(qn("w:sz"), tamanho)
        borda.set(qn("w:color"), cor)


def configurar_estilos(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(MARROM)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for nome, tamanho, cor, antes, depois in (
        ("Heading 1", 17, MARROM, 16, 7),
        ("Heading 2", 13, CARAMELO, 11, 5),
        ("Heading 3", 11, MARROM, 8, 4),
    ):
        estilo = doc.styles[nome]
        estilo.font.name = "Aptos Display"
        estilo._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        estilo._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        estilo.font.size = Pt(tamanho)
        estilo.font.bold = True
        estilo.font.color.rgb = RGBColor.from_string(cor)
        estilo.paragraph_format.space_before = Pt(antes)
        estilo.paragraph_format.space_after = Pt(depois)
        estilo.paragraph_format.keep_with_next = True


def adicionar_titulo(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(38)
    p.paragraph_format.space_after = Pt(4)
    definir_fonte(p.add_run("APPONO"), "Aptos Display", 13, CARAMELO, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    definir_fonte(p.add_run("Roteiro de Apresentação do MVP"), "Aptos Display", 27, MARROM, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    definir_fonte(
        p.add_run("AC2 - Projeto Inspira Startup | Duração planejada: 11 minutos"),
        tamanho=11,
        cor=CINZA,
        italico=True,
    )

    adicionar_caixa(
        doc,
        "Objetivo da apresentação",
        "Demonstrar o fluxo funcional da Appono: cadastro, autenticação, validações com APIs, "
        "persistência no Supabase, reserva automática de mesa e pedido antecipado vinculado à reserva.",
        DOURADO,
    )


def adicionar_caixa(doc, titulo, texto, cor_faixa=DOURADO):
    tabela = doc.add_table(rows=1, cols=2)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.autofit = False
    tabela.columns[0].width = Inches(0.12)
    tabela.columns[1].width = Inches(6.18)
    tabela.cell(0, 0).width = Inches(0.12)
    tabela.cell(0, 1).width = Inches(6.18)
    sombrear_celula(tabela.cell(0, 0), cor_faixa)
    sombrear_celula(tabela.cell(0, 1), CREME)
    margens_celula(tabela.cell(0, 0), 0, 0, 0, 0)
    margens_celula(tabela.cell(0, 1), 150, 180, 150, 180)
    p = tabela.cell(0, 1).paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    definir_fonte(p.add_run(titulo + "\n"), tamanho=10.5, cor=MARROM, negrito=True)
    definir_fonte(p.add_run(texto), tamanho=10, cor=CINZA)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def adicionar_linha_tempo(doc):
    doc.add_heading("Cronograma geral", level=1)
    tabela = doc.add_table(rows=1, cols=4)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.autofit = False
    larguras = [1.05, 1.25, 2.35, 1.85]
    cabecalhos = ["Tempo", "Responsável", "Parte", "Demonstração"]
    for i, (cabecalho, largura) in enumerate(zip(cabecalhos, larguras)):
        celula = tabela.rows[0].cells[i]
        celula.width = Inches(largura)
        sombrear_celula(celula, MARROM)
        margens_celula(celula)
        celula.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = celula.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        definir_fonte(p.add_run(cabecalho), tamanho=9.5, cor=BRANCO, negrito=True)

    linhas = [
        ("0:00-0:40", "Pessoa 1", "Problema e proposta", "Tela inicial"),
        ("0:40-3:30", "Pessoa 1", "Cadastro, login e acesso", "Cadastro de cliente + login"),
        ("3:30-6:30", "Pessoa 2", "APIs e cadastro empresarial", "ViaCEP + ReceitaWS"),
        ("6:30-10:20", "Pessoa 3", "Reserva e pedido antecipado", "Cliente + restaurante"),
        ("10:20-11:00", "Pessoa 3", "Conclusão e próximos passos", "GitHub e fechamento"),
        ("11:00-12:00", "Margem", "Imprevistos ou perguntas", "Não planejar conteúdo novo"),
    ]
    for indice, linha in enumerate(linhas):
        celulas = tabela.add_row().cells
        for i, (valor, largura) in enumerate(zip(linha, larguras)):
            celulas[i].width = Inches(largura)
            sombrear_celula(celulas[i], CREME if indice % 2 == 0 else "FFFFFF")
            margens_celula(celulas[i])
            celulas[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = celulas[i].paragraphs[0]
            if i < 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            definir_fonte(p.add_run(valor), tamanho=9.2, cor=MARROM, negrito=(i == 0))
    bordas_tabela(tabela)


def adicionar_fala(doc, texto):
    tabela = doc.add_table(rows=1, cols=1)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.autofit = False
    tabela.cell(0, 0).width = Inches(6.3)
    sombrear_celula(tabela.cell(0, 0), CREME)
    margens_celula(tabela.cell(0, 0), 140, 180, 140, 180)
    p = tabela.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    definir_fonte(p.add_run('“' + texto + '”'), tamanho=10.2, cor=MARROM, italico=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def adicionar_bullets(doc, itens):
    for item in itens:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(3)
        definir_fonte(p.add_run(item), tamanho=10.2)


def adicionar_passos(doc, itens):
    for numero, item in enumerate(itens, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.32)
        p.paragraph_format.first_line_indent = Inches(-0.32)
        p.paragraph_format.space_after = Pt(3)
        definir_fonte(p.add_run(f"{numero}. "), tamanho=10.2, cor=CARAMELO, negrito=True)
        definir_fonte(p.add_run(item), tamanho=10.2)


def adicionar_codigo(doc, caminho, codigo, explicacoes):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    definir_fonte(p.add_run("Arquivo: "), tamanho=9.5, cor=CINZA, negrito=True)
    definir_fonte(p.add_run(caminho), nome="Consolas", tamanho=9, cor=CARAMELO)

    tabela = doc.add_table(rows=1, cols=1)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.autofit = False
    tabela.cell(0, 0).width = Inches(6.3)
    sombrear_celula(tabela.cell(0, 0), "27211E")
    margens_celula(tabela.cell(0, 0), 140, 160, 140, 160)
    p = tabela.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    for indice, linha in enumerate(codigo.splitlines()):
        definir_fonte(p.add_run(linha), nome="Consolas", tamanho=8.3, cor="F7F0E7")
        if indice < len(codigo.splitlines()) - 1:
            p.add_run("\n")
    adicionar_bullets(doc, explicacoes)


def adicionar_integrante_1(doc):
    doc.add_page_break()
    doc.add_heading("Pessoa 1 - Proposta, Cadastro e Autenticação", level=1)
    adicionar_caixa(doc, "Tempo-alvo: 3 minutos e 30 segundos", "Apresentar o propósito da startup e demonstrar o acesso seguro ao sistema.", CARAMELO)

    doc.add_heading("1. Abertura", level=2)
    adicionar_fala(
        doc,
        "Bom dia. Nosso projeto se chama Appono. A startup conecta clientes e restaurantes por meio "
        "de reservas com pedidos antecipados, permitindo que o cliente chegue no horário marcado e "
        "encontre seu pedido próximo de estar pronto."
    )
    adicionar_fala(
        doc,
        "A aplicação possui dois módulos: cliente e restaurante. O frontend foi desenvolvido em Next.js, "
        "o backend utiliza Node.js com Express, e os dados são armazenados no PostgreSQL do Supabase."
    )

    doc.add_heading("2. Demonstração", level=2)
    adicionar_passos(
        doc,
        [
            "Abrir o cadastro de cliente e destacar os campos obrigatórios.",
            "Mostrar a máscara do CPF e tentar prosseguir com um CPF inválido.",
            "Concluir um cadastro preparado previamente ou explicar que os dados serão persistidos.",
            "Realizar login com uma conta de teste.",
            "Mostrar o redirecionamento automático para o dashboard correto.",
            "Abrir Configurações e mostrar que os dados cadastrados aparecem no perfil.",
        ],
    )

    doc.add_heading("3. Explicação técnica curta", level=2)
    adicionar_fala(
        doc,
        "O Supabase Auth valida as credenciais e devolve um token de acesso. Nas rotas protegidas, "
        "o backend valida esse token antes de permitir consultas ou alterações."
    )
    adicionar_codigo(
        doc,
        "backend/src/routes/auth.ts",
        '''authRouter.post("/login", async (req, res) => {
  const { email, password } = req.body;
  const { data, error } =
    await supabaseAuth.auth.signInWithPassword({ email, password });
});''',
        [
            "A rota POST recebe os dados enviados pelo formulário de login.",
            "req.body contém o e-mail e a senha informados pelo usuário.",
            "await aguarda o Supabase validar as credenciais.",
            "data recebe a sessão; error informa uma falha de autenticação.",
        ],
    )
    adicionar_codigo(
        doc,
        "backend/src/middleware/auth.ts",
        '''const accessToken = authorization.slice(7);
const { data: { user }, error } =
  await supabaseAuth.auth.getUser(accessToken);''',
        [
            "slice(7) remove o prefixo Bearer e preserva somente o token.",
            "getUser confirma se o token é válido antes de liberar a rota.",
        ],
    )


def adicionar_integrante_2(doc):
    doc.add_page_break()
    doc.add_heading("Pessoa 2 - APIs, Máscaras e Validações", level=1)
    adicionar_caixa(doc, "Tempo-alvo: 3 minutos", "Demonstrar duas APIs reais e explicar a integração pelo backend.", CARAMELO)

    doc.add_heading("1. Demonstração das APIs", level=2)
    adicionar_passos(
        doc,
        [
            "Abrir o cadastro de restaurante.",
            "Digitar um CEP válido e mostrar o preenchimento automático do endereço.",
            "Digitar um CNPJ válido e mostrar razão social, nome fantasia e situação retornados.",
            "Mostrar as máscaras de CEP e CNPJ.",
            "Abrir Configurações do restaurante e mostrar os dados persistidos.",
            "Explicar que CNPJ e razão social são imutáveis por segurança cadastral.",
        ],
    )

    doc.add_heading("2. Fala sugerida", level=2)
    adicionar_fala(
        doc,
        "Utilizamos a API ViaCEP para consultar o endereço e a ReceitaWS para consultar informações "
        "públicas do CNPJ. As chamadas passam pelo nosso backend, onde tratamos erros, definimos tempo "
        "limite e usamos cache para evitar consultas repetidas."
    )
    adicionar_fala(
        doc,
        "O CPF é validado matematicamente pelo backend, utilizando os dígitos verificadores. "
        "Assim, a validação não depende apenas da máscara visual do formulário."
    )

    doc.add_heading("3. Explicação de código", level=2)
    adicionar_codigo(
        doc,
        "backend/src/services/validacoes/cep.ts",
        '''const url = `https://viacep.com.br/ws/${cep}/json/`;
const resposta = await fetch(url, {
  headers: { Accept: "application/json" },
  signal: AbortSignal.timeout(TEMPO_LIMITE_REQUISICAO_MS),
});''',
        [
            "A URL inclui o CEP informado pelo usuário.",
            "fetch realiza a requisição para a API externa.",
            "Accept informa que a aplicação espera receber JSON.",
            "O timeout evita que a aplicação fique aguardando indefinidamente.",
        ],
    )
    adicionar_codigo(
        doc,
        "backend/src/services/validacoes/cpf.ts",
        '''if (cpf.length !== 11 || possuiDigitosRepetidos(cpf)) {
  return false;
}
return cpf.endsWith(`${primeiroDigito}${segundoDigito}`);''',
        [
            "A primeira condição rejeita tamanho incorreto e números repetidos.",
            "Os dígitos verificadores calculados são comparados com o final do CPF.",
        ],
    )


def adicionar_integrante_3(doc):
    doc.add_page_break()
    doc.add_heading("Pessoa 3 - Reserva e Pedido Antecipado", level=1)
    adicionar_caixa(doc, "Tempo-alvo: 3 minutos e 50 segundos", "Demonstrar o principal diferencial funcional da Appono.", CARAMELO)

    doc.add_heading("1. Demonstração do cliente", level=2)
    adicionar_passos(
        doc,
        [
            "Entrar no módulo do cliente e escolher um restaurante.",
            "Selecionar data, horário e quantidade de pessoas.",
            "Mostrar o consumo mínimo calculado para a reserva.",
            "Confirmar a reserva e abrir a página Meus Agendamentos.",
            "Explicar que o banco seleciona automaticamente uma mesa disponível.",
            "Mostrar a opção de pedido antecipado vinculada à reserva.",
        ],
    )

    doc.add_heading("2. Demonstração do restaurante", level=2)
    adicionar_passos(
        doc,
        [
            "Entrar no módulo do restaurante.",
            "Abrir Reservas e mostrar cliente, horário, pessoas e mesa selecionada.",
            "Mostrar o pedido antecipado relacionado à reserva, quando houver itens cadastrados.",
            "Explicar os status do pedido: confirmado, em preparo, pronto e entregue.",
        ],
    )

    doc.add_heading("3. Fala sugerida", level=2)
    adicionar_fala(
        doc,
        "Ao confirmar a reserva, uma função do PostgreSQL procura uma mesa disponível com capacidade "
        "suficiente. Essa verificação ocorre diretamente no banco para evitar que duas pessoas reservem "
        "a mesma mesa no mesmo horário."
    )
    adicionar_fala(
        doc,
        "O pedido antecipado é vinculado à reserva, ao cliente e ao restaurante. O horário da reserva "
        "serve como prazo operacional para que a cozinha organize o preparo."
    )

    doc.add_heading("4. Explicação de código", level=2)
    adicionar_codigo(
        doc,
        "backend/src/routes/reservations.ts",
        '''const { data, error } = await supabase.rpc(
  "criar_reserva_com_mesa_disponivel",
  {
    restaurante_id: body.id_restaurante,
    data_escolhida: body.data_reserva,
    inicio: body.horario_inicio,
    fim: body.horario_fim,
    pessoas: body.quantidade_pessoas,
  },
);''',
        [
            "rpc executa uma função criada diretamente no PostgreSQL.",
            "A função recebe restaurante, data, intervalo e quantidade de pessoas.",
            "O banco seleciona uma mesa compatível e bloqueia conflitos de horário.",
        ],
    )
    adicionar_codigo(
        doc,
        "backend/src/routes/orders.ts",
        '''const { data, error } = await supabase.rpc(
  "criar_pedido_antecipado",
  {
    reserva_id: body.id_reserva,
    itens: body.itens,
    observacoes_cliente: body.observacoes ?? null,
  },
);''',
        [
            "reserva_id cria a relação direta entre pedido e reserva.",
            "Os itens são enviados juntos para uma operação atômica.",
            "Se qualquer item for inválido, nenhum pedido incompleto é salvo.",
        ],
    )


def adicionar_perguntas(doc):
    doc.add_page_break()
    doc.add_heading("Perguntas Prováveis da Banca", level=1)
    perguntas = [
        ("Por que usar Supabase?", "Ele oferece PostgreSQL, autenticação, armazenamento e regras de segurança integradas, acelerando o desenvolvimento do MVP."),
        ("O que é frontend?", "É a interface visual utilizada pelo usuário. No projeto, foi desenvolvida com Next.js."),
        ("O que é backend?", "É o servidor Node.js com Express responsável por regras, validações e comunicação segura com o banco."),
        ("O que é uma API?", "É uma interface de comunicação entre sistemas por meio de requisições e respostas."),
        ("O que é async/await?", "É uma forma de aguardar operações assíncronas, como consultas ao banco e APIs externas, sem bloquear toda a aplicação."),
        ("O que é uma migration?", "É um arquivo versionado que registra alterações estruturais no banco e permite reproduzi-las em outros ambientes."),
        ("O que é uma chave estrangeira?", "É uma regra que conecta tabelas. Por exemplo, id_reserva conecta um pedido à reserva correspondente."),
        ("O que é RLS?", "Row Level Security são regras do Supabase que controlam quais registros cada usuário pode consultar ou alterar."),
        ("Por que consultar APIs pelo backend?", "Para centralizar validações, tratar erros, usar cache e evitar confiar somente no navegador."),
        ("Como o sistema evita duas reservas para a mesma mesa?", "A disponibilidade é validada dentro do PostgreSQL, que possui uma restrição contra sobreposição de reservas ativas."),
    ]
    tabela = doc.add_table(rows=1, cols=2)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.autofit = False
    larguras = [2.2, 4.1]
    for i, titulo in enumerate(("Pergunta", "Resposta recomendada")):
        tabela.rows[0].cells[i].width = Inches(larguras[i])
        sombrear_celula(tabela.rows[0].cells[i], MARROM)
        margens_celula(tabela.rows[0].cells[i])
        p = tabela.rows[0].cells[i].paragraphs[0]
        definir_fonte(p.add_run(titulo), tamanho=9.5, cor=BRANCO, negrito=True)
    for indice, (pergunta, resposta) in enumerate(perguntas):
        celulas = tabela.add_row().cells
        for i, texto in enumerate((pergunta, resposta)):
            celulas[i].width = Inches(larguras[i])
            sombrear_celula(celulas[i], CREME if indice % 2 == 0 else "FFFFFF")
            margens_celula(celulas[i], 110, 140, 110, 140)
            celulas[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = celulas[i].paragraphs[0]
            definir_fonte(p.add_run(texto), tamanho=9.2, cor=MARROM, negrito=(i == 0))
    bordas_tabela(tabela)


def adicionar_checklist(doc):
    doc.add_page_break()
    doc.add_heading("Checklist Final da Equipe", level=1)
    adicionar_caixa(
        doc,
        "Regra principal",
        "Demonstrem somente fluxos testados. Caso algo falhe, expliquem o comportamento esperado e avancem para a próxima etapa.",
        VERMELHO,
    )
    doc.add_heading("Antes da apresentação", level=2)
    adicionar_bullets(
        doc,
        [
            "Executar npm run dev e confirmar que frontend e backend iniciaram sem erros.",
            "Testar login de uma conta cliente e de uma conta restaurante.",
            "Preparar dados válidos para CEP e CNPJ.",
            "Manter uma reserva confirmada previamente para contingência.",
            "Confirmar que o Supabase está acessível.",
            "Preparar o link do GitHub e conferir se arquivos .env não foram enviados.",
            "Fechar notificações e abas desnecessárias.",
        ],
    )
    doc.add_heading("Durante a apresentação", level=2)
    adicionar_bullets(
        doc,
        [
            "Uma pessoa opera o computador enquanto outra fala.",
            "Não ler todos os textos das telas; explicar o valor e demonstrar a ação.",
            "Cada integrante deve encerrar sua parte conectando-a à próxima.",
            "Evitar abrir arquivos .env, chaves do Supabase ou dados sensíveis.",
            "Usar no máximo 11 minutos de conteúdo e preservar a margem final.",
        ],
    )
    doc.add_heading("Transições entre integrantes", level=2)
    adicionar_fala(doc, "Pessoa 1 para Pessoa 2: Agora que o usuário está autenticado, vamos mostrar como garantimos a qualidade dos dados cadastrados usando validações e APIs externas.")
    adicionar_fala(doc, "Pessoa 2 para Pessoa 3: Com o restaurante validado e cadastrado, podemos demonstrar o principal fluxo da plataforma: a reserva e o pedido antecipado.")


def adicionar_rodape(doc):
    for secao in doc.sections:
        rodape = secao.footer
        p = rodape.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        definir_fonte(p.add_run("APPONO | Roteiro de apresentação do MVP | 11/06/2026"), tamanho=8.5, cor=CINZA)


def gerar():
    SAIDA.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    secao = doc.sections[0]
    secao.top_margin = Inches(0.72)
    secao.bottom_margin = Inches(0.72)
    secao.left_margin = Inches(0.9)
    secao.right_margin = Inches(0.9)
    secao.header_distance = Inches(0.35)
    secao.footer_distance = Inches(0.35)

    configurar_estilos(doc)
    adicionar_titulo(doc)
    adicionar_linha_tempo(doc)
    adicionar_integrante_1(doc)
    adicionar_integrante_2(doc)
    adicionar_integrante_3(doc)
    adicionar_perguntas(doc)
    adicionar_checklist(doc)
    adicionar_rodape(doc)
    doc.save(SAIDA)
    print(SAIDA)


if __name__ == "__main__":
    gerar()
